import os
import re
import json
import time
import contextlib
import io
import random
from pathlib import Path

import keyboard
import pyperclip
from winotify import Notification

from docx import Document
from docx2pdf import convert

from openai import OpenAI, OpenAIError


# ================= CONFIG ================= #

MODEL = "gpt-4.1-mini"
TEMPERATURE_MAIN = 0.35
TEMPERATURE_REGEN = 0.70

COMPANY_COUNT = 4

# IMPORTANT: You generate 4 people, so reflect that everywhere
PERSON_ORDER = ["Timothy", "Wilfredo", "Lou", "Ryan"]

TEMPLATES = {
    "Timothy": "templates/Timothy.docx",
    "Wilfredo": "templates/Wilfredo.docx",
    "Lou": "templates/Lou.docx",
    "Ryan": "templates/Ryan.docx",
}

SKILL_CATEGORY_ORDER = [
    "Programming Languages",
    "Frameworks & Libraries",
    "Databases",
    "Cloud & DevOps",
    "Testing",
    "Tools",
]

OUTPUT_DIR = "output"
BULLET = "  \u2022    "  # bullet point prefix

# ================= PERSON PROFILES ================= #

PERSON_PROFILES = {
    "Timothy": {
        "style": "strategic, architecture-focused, leadership-driven",
        "summary": "senior software engineer focused on scalable systems",
        "metrics": "latency, uptime, scale, cost, reliability"
    },
    "Wilfredo": {
        "style": "hands-on, delivery-focused, optimization-oriented",
        "summary": "experienced senior full-stack engineer with strong execution",
        "metrics": "throughput, build time, cost, delivery speed, defect reduction"
    },
    "Lou": {
        "style": "collaborative, balanced, system improvement oriented",
        "summary": "senior software engineer with growing ownership and impact",
        "metrics": "MTTR, incident rate, test coverage, availability, customer satisfaction"
    },
    "Ryan": {
        "style": "practical, implementation-heavy, learning-focused",
        "summary": "senior software engineer with solid fundamentals and growth mindset",
        "metrics": "feature adoption, support tickets, time-to-ship, performance, usage"
    },
}


# ================= DEFAULT SKILLS ================= #

DEFAULT_SKILLS = {
    "Programming Languages": ["JavaScript", "TypeScript", "Python", "PHP", "Ruby", "Golang", "Java", "C#"],
    "Frameworks & Libraries": ["Microservices", "REST APIs", "React.js", "Next.js", "Node.js", "Django", "Flask", "Laravel"],
    "Databases": ["MySQL", "PostgreSQL", "MongoDB", "Firebase", "Redis"],
    "Testing": ["Jest", "React Testing Library", "Cypress", "Mocha", "Go-Testify"],
    "Cloud & DevOps": ["AWS", "Azure", "Docker", "Jenkins"],
    "Tools": ["Git", "Jira", "CI/CD"]
}

MAX_SKILLS_PER_CATEGORY = 8


# ================= INIT ================= #

def notify(message: str):
    with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
        Notification(
            app_id="Resume Generator",
            title="Resume Generator",
            msg=message,
            duration="short"
        ).show()


def get_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("Missing OPENAI_API_KEY environment variable.")
    return OpenAI(api_key=api_key)



client = get_client()

def call_openai_json(
    prompt: str,
    model: str = "gpt-4.1-mini",
    temperature: float = 0.3,
    max_retries: int = 4,
    retry_sleep: float = 2.0,
) -> dict:
    """
    Calls OpenAI and guarantees a parsed JSON object or raises.
    Designed for ATS / resume generation (strict JSON).
    """

    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            response = client.responses.create(
                model=model,
                input=prompt,
                temperature=temperature,
                timeout=90,
            )

            raw = extract_text(response)
            if not raw:
                raise ValueError("Empty response")

            # --- Fast path: valid JSON ---
            try:
                return json.loads(raw)
            except json.JSONDecodeError:
                pass

            # --- Recovery path: extract first JSON object ---
            match = re.search(r"\{[\s\S]*\}", raw)
            if match:
                candidate = match.group(0)
                try:
                    return json.loads(candidate)
                except json.JSONDecodeError:
                    pass

            raise ValueError("Model returned invalid JSON")

        except OpenAIError as e:
            last_error = e
            # Rate-limit or transient error → retry
            if "rate limit" in str(e).lower() or "timeout" in str(e).lower():
                time.sleep(retry_sleep * attempt)
                continue
            raise

        except Exception as e:
            last_error = e
            time.sleep(retry_sleep * attempt)

    raise RuntimeError(f"OpenAI JSON call failed after {max_retries} attempts: {last_error}")

output_base_dir = None
folder_opened = False
is_running = False


# ================= PROMPTS ================= #

def build_prompt_compress_jd(jd_text: str) -> str:
    return f"""
Return ONLY valid JSON.

TASK:
Compress this Job Description into an ATS_PACKAGE JSON object with:
- job_title_exact (the exact job title phrase in the JD if present; else "Senior Software Engineer")
- core_hard: list of hard skills/tech/phrases that must be present (20–35 max), something like 
- core_soft: list of soft skills phrases verbatim from JD (10–20 max)
- required_phrases: important long phrases (5–15 max)
- constraints: {{"must_appear_each_resume": true, "max_repeat_per_term": 3}}
- lane_spines: 5 distinct focus spines for Timothy/Wilfredo/Lou/Ryan (1 line each)

RULES:
- Use phrases verbatim from JD when possible.
- Do NOT invent technologies not in JD.
- Output JSON only.

JOB DESCRIPTION:
{jd_text}
""".strip()

def build_avoid_profile(resume: dict) -> dict:
    bullets = []
    for c in ["company_1","company_2","company_3","company_4"]:
        bullets += resume.get("experience", {}).get(c, []) or []

    def first5(s):
        w = re.findall(r"[A-Za-z0-9]+", s)
        return " ".join(w[:5]).lower()

    starts = sorted({first5(b) for b in bullets if isinstance(b,str)})

    verbs = []
    for b in bullets:
        m = re.match(r"^\s*([A-Za-z]+)", b.strip())
        if m:
            verbs.append(m.group(1).lower())

    return {
        "banned_starts_first5": starts[:80],
        "verbs_used": sorted(set(verbs))[:60],
    }

import json

def build_prompt_one_person_from_pkg(
    ats_package: dict,
    person: str,
    avoid_profile: dict,
    secondary_json: str,
) -> str:
    profile = PERSON_PROFILES[person]
    
    pkg = json.dumps(ats_package, ensure_ascii=False)
    avoid = json.dumps(avoid_profile, ensure_ascii=False)

    return f"""
Output ONLY valid JSON (no markdown, no extra text).

ATS_PACKAGE (source of truth):
{pkg}

PERSON PROFILE (guidance only — DO NOT quote directly):
- Writing style: {profile["style"]}
- Resume positioning: {profile["summary"]}
- Preferred metric families: {profile["metrics"]}

AVOID_PROFILE (must not match, hard uniqueness constraints):
{avoid}

SECONDARY TECH (optional, use naturally, do NOT replace core skills):
{secondary_json}

TASK:
Generate ONE senior-level resume for "{person}" using ATS_PACKAGE.

HARD RULES:
- DO NOT repeat core_soft terms more than 3 times (critical!).
- ALL terms in ATS_PACKAGE.core_hard + core_soft + required_phrases MUST appear at least once in THIS resume(critical!).
- Do not repeat any term more than ATS_PACKAGE.constraints.max_repeat_per_term times.
- Do NOT start any bullet with any string in AVOID_PROFILE.banned_starts_first5.
- Avoid reusing verbs in AVOID_PROFILE.verbs_used as much as possible (prefer new verbs).
- Follow the focus spine in ATS_PACKAGE.lane_spines["{person}"].
- Follow the PERSON PROFILE to choose focus, verbs, and metrics.
- if core_soft is less than 5, add some soft skills.
- If core_hard is less than 5, add some hard skills.

OUTPUT JSON SCHEMA:
{{
  "{person}": {{
    "summary": "string",
    "skills": {{
      "Programming Languages": [],
      "Frameworks & Libraries": [],
      "Databases": [],
      "Cloud & DevOps": [],
      "Testing": [],
      "Tools": []
    }},
    "experience": {{
      "company_1": [],
      "company_2": [],
      "company_3": [],
      "company_4": []
    }}
  }}
}}

EXPERIENCE RULES:
- company_1 and company_2 → Senior-level responsibilities, 6–8 bullets each
- company_3 → Mid-level responsibilities, 5 bullets
- company_4 → Junior-level responsibilities, 4–5 bullets
- Sometimes include core_soft skill (verbatim)
- EACH company section MUST:
    - Include at least TWO core_soft skill BUT don't repeat any (verbatim)
- EACH bullet MUST:
  - Include at least ONE JD hard skill (verbatim)
  - Include a measurable metric (%,$,users,latency,scale)
- Enforce uniqueness : different verbs, different metric types, different framing.

QUALITY:
- Summary starts with "Senior Software Engineer"
- Each bullet includes: 1 core_hard term + (more than 0.2 core_soft phrase) + 1 metric
- Make bullets structurally varied (metric-first, collaboration-first, method-first, outcome-first).

Return JSON only.
""".strip()

# ================= CHOOSE RANDOM SKILLS ================= #

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())

def pick_random_secondary_skills(
    jd_text: str,
    default_skills: dict,
    per_category_min: int = 0,
    per_category_max: int = 2,
    total_max: int = 8,
    seed: int | None = None,
) -> dict:
    """
    Randomly pick secondary skills per category, but keep ATS-safe:
    - Avoid conflicting stacks (e.g., Angular JD => avoid React/Next)
    - Avoid adding too many irrelevant frameworks/tools
    """
    rng = random.Random(seed)
    jd = _norm(jd_text)

    has_angular = "angular" in jd
    has_java = "java" in jd

    # Conflicts / avoid list based on JD
    banned = set()
    if has_angular:
        banned.update({_norm("React.js"), _norm("Next.js")})
    # If JD is strongly Java/Angular enterprise, you may also avoid Laravel/Ruby etc.
    if has_java:
        banned.update({_norm("Laravel"), _norm("Ruby")})

    picked = {cat: [] for cat in default_skills.keys()}

    # Random pick by category
    for cat, items in default_skills.items():
        pool = [x for x in items if _norm(x) not in banned]
        if not pool:
            continue

        k = rng.randint(per_category_min, min(per_category_max, len(pool)))
        chosen = rng.sample(pool, k) if k > 0 else []
        picked[cat] = chosen

    # Enforce total_max across categories
    flat = [(cat, s) for cat, lst in picked.items() for s in lst]
    rng.shuffle(flat)
    flat = flat[:total_max]

    trimmed = {cat: [] for cat in default_skills.keys()}
    for cat, s in flat:
        trimmed[cat].append(s)

    return trimmed

def build_secondary_by_person(jd_text: str) -> dict:
    secondary_by_person = {}
    for person in PERSON_ORDER:
        # different seed per person for stable randomness
        seed = abs(hash(person)) % (10**9)
        secondary_by_person[person] = pick_random_secondary_skills(
            jd_text=jd_text,
            default_skills=DEFAULT_SKILLS,
            per_category_min=0,
            per_category_max=2,
            total_max=8,
            seed=seed,
        )
    return secondary_by_person


# ================= OPENAI RESPONSE PARSING ================= #

def extract_text(response) -> str:
    if hasattr(response, "output_text") and response.output_text:
        return response.output_text.strip()

    texts = []
    for message in getattr(response, "output", []):
        for block in getattr(message, "content", []):
            t = getattr(block, "text", None)
            if t:
                texts.append(t)
    return "\n".join(texts).strip()


# ================= SKILLS MERGE ================= #

def merge_skills(gpt_skills: dict) -> dict:
    merged = {}
    for category in set(gpt_skills) | set(DEFAULT_SKILLS):
        s = set(gpt_skills.get(category, []) or [])
        s.update(DEFAULT_SKILLS.get(category, []) or [])
        merged[category] = sorted(s)[:MAX_SKILLS_PER_CATEGORY]
    return merged


# ================= DOCX HELPERS ================= #

def replace_placeholder_plain(doc: Document, placeholder: str, value: str):
    # NOTE: This destroys runs/formatting in that paragraph. Only use for plain sections.
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, value)

import random
from typing import List, Any

def shuffle_and_drop(items: List[Any], max_drop: int = 2) -> List[Any]:
    """
    Shuffles the list and removes 0–2 elements.
    When list length < 4, removal is biased toward 0 or 1.
    """
    if not items:
        return []

    result = items.copy()
    random.shuffle(result)
    n = len(result)

    # Drop logic with intelligent bias
    if n < 2:
        drop_count = 0
    elif n < 4:
        # Strong bias toward 0 or 1
        drop_count = random.choices(
            population=[0, 1, 2],
            weights=[0.6, 0.35, 0.05],
            k=1
        )[0]
    else:
        # Normal behavior for larger lists
        drop_count = random.randint(0, 2)

    # Safety clamp
    drop_count = min(drop_count, n)

    # Remove random elements
    for _ in range(drop_count):
        result.pop(random.randrange(len(result)))

    return result


def replace_skills_placeholder(doc: Document, placeholder: str, skills: dict):
    # Bold only categories using runs
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.clear()
            for category in SKILL_CATEGORY_ORDER:
                items = skills.get(category, [])
                if not items:
                    continue
                run_cat = p.add_run(f"{category}: ")
                run_cat.bold = True
                reorderd_items = shuffle_and_drop(items)
                p.add_run(", ".join(reorderd_items))
                p.add_run("\n")


def fill_template(template_path: str, data: dict, output_path: str):
    doc = Document(template_path)

    replace_placeholder_plain(doc, "{{SUMMARY}}", data.get("summary", ""))

    merged_skills = merge_skills(data.get("skills", {}))
    replace_skills_placeholder(doc, "{{SKILLS}}", merged_skills)

    for i in range(1, COMPANY_COUNT + 1):
        bullets_list = data.get("experience", {}).get(f"company_{i}", []) or []
        bullets = f"\n{BULLET}".join(bullets_list)
        bullets = bullets.removesuffix(BULLET)
        bullets = BULLET + bullets if bullets else ""
        replace_placeholder_plain(doc, f"{{{{EXP_COMPANY_{i}}}}}", bullets)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def batch_convert_to_pdf(folder: str):
    folder_path = Path(folder)
    for docx_path in folder_path.glob("*.docx"):
        pdf_path = folder_path / f"{docx_path.stem}.pdf"
        convert(str(docx_path), str(pdf_path))


# ================= MAIN ================= #
def on_hotkey_generate():
    """
    New pipeline:
    1) Compress JD once -> ATS_PACKAGE (JSON)
    2) Generate 5 resumes sequentially (one person at a time)
       using ATS_PACKAGE + growing AVOID_PROFILE (no repeats)
    3) Write DOCX + convert to PDF
    """
    global output_base_dir, folder_opened, is_running

    if is_running:
        return
    is_running = True

    try:
        jd = pyperclip.paste()
        if not jd.strip():
            notify("Clipboard is empty")
            return

        if output_base_dir is None:
            output_base_dir = os.path.join(OUTPUT_DIR, "JD_resumes")
            os.makedirs(output_base_dir, exist_ok=True)

        # ---------- 1) Compress JD once ----------
        notify("Compressing JD (ATS package)...")
        compress_prompt = build_prompt_compress_jd(jd)
        
        ats_package = call_openai_json(compress_prompt)  # must return dict
        print(ats_package)
        # ---------- 2) Generate resumes sequentially ----------
        notify("Generating 4 resumes (sequential, low tokens)...")

        all_data = {}
        avoid_profile = {
            "banned_starts_first5": [],
            "verbs_used": [],
        }
        secondary_by_person = build_secondary_by_person(jd)
        

        for person in PERSON_ORDER:
            secondary_for_person = secondary_by_person[person]
            secondary_json = json.dumps(secondary_for_person, ensure_ascii=False)

            notify(f"Generating {person} resume...")

            person_prompt = build_prompt_one_person_from_pkg(
                ats_package=ats_package,
                person=person,
                avoid_profile=avoid_profile,
                secondary_json=secondary_json,
            )

            one_json = call_openai_json(person_prompt)  # returns {"Timothy": {...}}
            if person not in one_json:
                raise ValueError(f"Missing '{person}' key in model output")

            pdata = one_json[person]
            all_data[person] = pdata

            # Update avoid profile to reduce similarity for next resumes
            ap = build_avoid_profile(pdata)
            avoid_profile["banned_starts_first5"] = sorted(
                set(avoid_profile["banned_starts_first5"] + ap.get("banned_starts_first5", []))
            )[:250]
            avoid_profile["verbs_used"] = sorted(
                set(avoid_profile["verbs_used"] + ap.get("verbs_used", []))
            )[:150]

            if not pdata:
                notify(f"{person} missing from generated data")
                continue
            
            template = TEMPLATES[person]
            output_docx = os.path.join(output_base_dir, f"{person}_resume.docx")
            fill_template(template, pdata, output_docx)
            notify(f"{person} DOCX generated")

            if not folder_opened:
                try:
                    os.startfile(output_base_dir)
                except Exception:
                    pass
                folder_opened = True

        # ---------- 4) Convert to PDF ----------
        notify("DOCX generation done. Converting to PDF...")
        batch_convert_to_pdf(output_base_dir)
        notify("All PDFs generated")

    except Exception as e:
        print("Error:", e)
        notify(f"Error generating resumes: {e}")
    finally:
        is_running = False

def on_hotkey_pdf_only():
    global output_base_dir, is_running

    if is_running:
        return
    is_running = True

    try:
        output_base_dir = os.path.join(OUTPUT_DIR, "JD_resumes")
        if not os.path.isdir(output_base_dir):
            notify("No output folder found. Generate resumes first.")
            return

        notify("Converting DOCX to PDF...")
        batch_convert_to_pdf(output_base_dir)
        notify("All PDFs generated")

    except Exception as e:
        print("Error:", e)
        notify("Error generating PDFs")
    finally:
        is_running = False


def main():
    notify("Resume Generator is running")
    keyboard.add_hotkey("ctrl+q", on_hotkey_generate)
    keyboard.add_hotkey("ctrl+alt+p", on_hotkey_pdf_only)
    keyboard.wait()


if __name__ == "__main__":
    main()
